import streamlit as st
import pandas as pd
import requests
import time
import json
import re
import zipfile
import base64
from io import BytesIO
from docx import Document
from huggingface_hub import InferenceClient

# Configure page
st.set_page_config(page_title="SEO Content Automation", page_icon="📚", layout="wide")

# Initialize session state
if "articles" not in st.session_state:
    st.session_state["articles"] = {}
if "images" not in st.session_state:
    st.session_state["images"] = {}
if "publish_log" not in st.session_state:
    st.session_state["publish_log"] = []

def init_hf_client():
    """Initialize Hugging Face client"""
    try:
        HF_TOKEN = st.secrets.get("HF_TOKEN")
        if not HF_TOKEN:
            return None
        return InferenceClient(
            model="stabilityai/stable-diffusion-3-medium",
            token=HF_TOKEN
        )
    except Exception as e:
        st.error(f"Error initializing HF client: {str(e)}")
        return None

def get_api_key(provider):
    """Get API key for selected provider"""
    if provider == "Grok (X.AI)":
        return st.secrets.get("GROK_API_KEY")
    elif provider == "OpenAI":
        return st.secrets.get("OPENAI_API_KEY")
    return None

def call_ai_for_metadata(keyword, intent, content_type, notes, api_key, provider):
    """Generate metadata using AI API"""
    prompt = f"""
You are a search volume estimator and content strategist for Indian students and professionals.

Given the keyword: "{keyword}"
Intent: {intent}
Content Type: {content_type}
Notes: {notes}

Tasks:
1. Estimate search volume as High, Medium, or Low based on Indian market
2. Create a high-CTR SEO title optimized for Indian audience (include year 2024/2025)
3. Generate a detailed outline with 5-7 bullet points including:
   - Introduction/What is section
   - Key features/benefits
   - Detailed information (eligibility, process, etc.)
   - Tables section (comparison/data)
   - FAQ section
   - Conclusion
4. Provide specific content instructions mentioning:
   - Target keyword usage frequency
   - Required tables (comparison, fees, eligibility, etc.)
   - FAQ requirements (5-8 questions)
   - Tone and style preferences
   - Indian context requirements

Respond in JSON format:
{{
  "volume": "High/Medium/Low",
  "seo_title": "Complete Guide to [Keyword] in India 2024 - Benefits, Process & FAQs",
  "outline": [
    "Introduction - What is [Keyword]?",
    "Key Features and Benefits of [Keyword]", 
    "Detailed [Keyword] Information and Requirements",
    "Comparison Table - [Keyword] Categories/Types",
    "Step-by-Step Application Process",
    "Eligibility Criteria and Documents Required",
    "Frequently Asked Questions (FAQs)",
    "Conclusion and Key Takeaways"
  ],
  "instructions": "Use keyword '[keyword]' 8-12 times naturally. Include 2 detailed tables: one for comparison/categories and one for fees/requirements. Add comprehensive FAQ section with 6-8 questions. Focus on Indian context with specific data. Use professional, informative tone. Include current 2024/2025 information."
}}
"""
    
    # Configure API based on provider
    if provider == "Grok (X.AI)":
        url = "https://api.x.ai/v1/chat/completions"
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        body = {
            "model": "grok-beta",
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.7
        }
    elif provider == "OpenAI":
        url = "https://api.openai.com/v1/chat/completions"
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        body = {
            "model": "gpt-4o-mini",
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.7
        }
    else:
        st.error(f"Unsupported provider: {provider}")
        return None

    try:
        response = requests.post(url, json=body, headers=headers)
        if response.status_code == 200:
            content = response.json()["choices"][0]["message"]["content"]
            try:
                json_data = json.loads(content)
                return json_data
            except:
                return {
                    "volume": "Medium",
                    "seo_title": f"Complete Guide to {keyword} in India 2024 - Benefits, Eligibility & Process",
                    "outline": [
                        f"What is {keyword}? - Introduction and Overview",
                        f"Key Features and Benefits of {keyword}",
                        f"Eligibility Criteria for {keyword}",
                        f"Application Process and Required Documents",
                        f"Detailed Information Table - {keyword} Categories",
                        f"Frequently Asked Questions (FAQs) about {keyword}",
                        "Conclusion and Important Points"
                    ],
                    "instructions": f"Use '{keyword}' naturally 10-15 times. Include 2 tables: eligibility criteria and comparison table. Add FAQ section with 6-8 questions. Focus on Indian context with current data. Professional tone."
                }
        else:
            st.error(f"API Error ({provider}): {response.status_code} - {response.text}")
            return None
    except Exception as e:
        st.error(f"Request failed ({provider}): {str(e)}")
        return None

def generate_article(keyword, seo_title, outline, instructions, api_key, provider):
    """Generate complete article based on metadata"""
    outline_text = "\n".join([f"- {point}" for point in outline])
    
    prompt = f"""
You are an expert content writer specializing in educational content for Indian students and professionals.

Write a complete SEO-optimized article for:
Keyword: "{keyword}"
Title: "{seo_title}"

Structure using this outline:
{outline_text}

Special Instructions:
{instructions}

CRITICAL FORMATTING REQUIREMENTS:
1. Use the target keyword "{keyword}" naturally throughout the article (aim for 1-2% keyword density)
2. Include keyword variations and related terms
3. Structure with clear H1, H2, H3 headings using HTML tags
4. Include at least 1-2 detailed tables with relevant data
5. Add a comprehensive FAQ section at the end with 5-8 questions
6. Use bullet points and numbered lists where appropriate
7. Include specific data, statistics, and numbers
8. Write in HTML format with proper semantic tags

Write the complete article now following this structure.
Requirements:
- 1000-1500 words
- Natural keyword usage throughout
- Include specific Indian context and data
- Use tables for complex information
- Comprehensive FAQ section
- Professional, informative tone
"""
    
    # Configure API based on provider
    if provider == "Grok (X.AI)":
        url = "https://api.x.ai/v1/chat/completions"
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        body = {
            "model": "grok-beta",
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.7,
            "max_tokens": 2000
        }
    elif provider == "OpenAI":
        url = "https://api.openai.com/v1/chat/completions"
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        body = {
            "model": "gpt-4o-mini",
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.7,
            "max_tokens": 2000
        }

    try:
        response = requests.post(url, json=body, headers=headers)
        if response.status_code == 200:
            article = response.json()["choices"][0]["message"]["content"]
            return article
        else:
            st.error(f"Article generation failed ({provider}): {response.status_code} - {response.text}")
            return None
    except Exception as e:
        st.error(f"Article generation error ({provider}): {str(e)}")
        return None

def apply_internal_links(article_content, anchor_map):
    """Apply internal links to article content"""
    linked_article = article_content
    
    # Replace first occurrence of each anchor (case-insensitive)
    for anchor, url in anchor_map.items():
        pattern = re.compile(rf"\b({re.escape(anchor)})\b", re.IGNORECASE)
        linked_article, n = pattern.subn(
            rf'<a href="{url}" target="_blank">\1</a>', 
            linked_article, 
            count=1
        )
    
    # Add related links table
    table_rows = "".join([
        f'<tr><td>{anchor}</td><td><a href="{url}" target="_blank">{url}</a></td></tr>'
        for anchor, url in anchor_map.items()
    ])
    
    link_table = """
<div style="margin-top: 30px;">
<h2>Related Links</h2>
<table border="1" cellpadding="8" cellspacing="0" style="border-collapse: collapse; width: 100%;">
<tr style="background-color: #f2f2f2;"><th>Topic</th><th>Link</th></tr>
""" + table_rows + """
</table>
</div>
"""
    
    return linked_article + link_table

def generate_ai_image(prompt, hf_client):
    """Generate image using Hugging Face Inference Client"""
    if not hf_client:
        st.error("Hugging Face client not initialized")
        return None
    
    try:
        # Generate image using the inference client
        image = hf_client.text_to_image(prompt)
        
        # Convert PIL image to BytesIO
        img_buffer = BytesIO()
        image.save(img_buffer, format='PNG')
        img_buffer.seek(0)
        
        return img_buffer
    except Exception as e:
        st.error(f"Image generation error: {str(e)}")
        return None

def publish_to_wordpress(keyword, content, image_buffer, tags, wp_config, publish_now=True):
    """Publish article to WordPress"""
    wp_base = wp_config["base_url"]
    auth_str = f"{wp_config['username']}:{wp_config['password']}"
    auth_token = base64.b64encode(auth_str.encode()).decode("utf-8")
    headers = {"Authorization": f"Basic {auth_token}"}
    
    img_id = None
    
    # Upload image if provided
    if image_buffer:
        try:
            image_buffer.seek(0)
            img_data = image_buffer.read()
            img_headers = headers.copy()
            img_headers.update({
                "Content-Disposition": f"attachment; filename={keyword.replace(' ', '_')}.jpg",
                "Content-Type": "image/jpeg"
            })
            media_url = f"{wp_base}/wp-json/wp/v2/media"
            img_resp = requests.post(media_url, headers=img_headers, data=img_data)
            
            if img_resp.status_code == 201:
                img_id = img_resp.json()["id"]
            else:
                st.warning(f"Image upload failed for {keyword}: {img_resp.text}")
        except Exception as e:
            st.error(f"Image upload error: {str(e)}")
    
    # Create/get tags
    tag_ids = []
    if tags:
        for tag in [t.strip() for t in tags.split(",") if t.strip()]:
            try:
                # Check if tag exists
                tag_check = requests.get(f"{wp_base}/wp-json/wp/v2/tags?search={tag}", headers=headers)
                if tag_check.status_code == 200 and tag_check.json():
                    tag_ids.append(tag_check.json()[0]["id"])
                else:
                    # Create new tag
                    tag_create = requests.post(f"{wp_base}/wp-json/wp/v2/tags", headers=headers, json={"name": tag})
                    if tag_create.status_code == 201:
                        tag_ids.append(tag_create.json()["id"])
            except Exception as e:
                st.warning(f"Tag creation failed for '{tag}': {str(e)}")
    
    # Publish article
    post_data = {
        "title": keyword,
        "content": content,
        "status": "publish" if publish_now else "draft",
        "tags": tag_ids
    }
    
    if img_id:
        post_data["featured_media"] = img_id
    
    try:
        post_resp = requests.post(f"{wp_base}/wp-json/wp/v2/posts", headers=headers, json=post_data)
        if post_resp.status_code == 201:
            post_url = post_resp.json()["link"]
            return {"success": True, "url": post_url}
        else:
            return {"success": False, "error": post_resp.text}
    except Exception as e:
        return {"success": False, "error": str(e)}

# Sidebar for API configuration
st.sidebar.header("🔧 API Configuration")

# AI Model Selection
ai_provider = st.sidebar.selectbox(
    "Choose AI Provider",
    ["Grok (X.AI)", "OpenAI"],
    help="Select your preferred AI provider for content generation"
)

# Get API key for selected provider
current_api_key = get_api_key(ai_provider)

# Initialize HF client
hf_client = init_hf_client()

# WordPress Config
st.sidebar.header("🌐 WordPress Settings")
wp_base_url = st.secrets.get("WP_BASE_URL", "")
wp_username = st.secrets.get("WP_USERNAME", "")
wp_password = st.secrets.get("WP_PASSWORD", "")

if wp_base_url and wp_username and wp_password:
    st.sidebar.success("✅ WordPress configured")
else:
    st.sidebar.warning("⚠️ WordPress not configured in secrets")

# Main App Interface
st.title("📚 SEO Content Automation Pipeline")
st.markdown("Upload topics → Generate metadata → Create articles → Add images → Publish to WordPress")

# Show current provider status
if current_api_key:
    st.success(f"✅ Connected to {ai_provider}")
else:
    st.error(f"❌ {ai_provider} API key not found in secrets")

# Tab interface
tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs([
    "📋 Topic Upload", 
    "📝 Article Generation", 
    "🔗 Internal Links", 
    "🖼️ Images", 
    "🚀 WordPress Publish", 
    "📊 Export & Logs"
])

with tab1:
    st.header("📋 Step 1: Topic Metadata Generation")
    
    uploaded_file = st.file_uploader("Upload Excel with keywords", type=["xlsx"])
    
    if uploaded_file and current_api_key:
        df = pd.read_excel(uploaded_file)
        
        required_cols = ["Keyword", "Intent", "Content Type", "Notes"]
        if not all(col in df.columns for col in required_cols):
            st.error(f"❌ Your Excel must contain columns: {', '.join(required_cols)}")
        else:
            st.success(f"✅ Found {len(df)} topics to process")
            
            if st.button("🚀 Generate Metadata for All Topics"):
                progress_bar = st.progress(0)
                results = []
                
                for i, row in df.iterrows():
                    st.info(f"Processing: {row['Keyword']} ({i+1}/{len(df)})")
                    
                    output = call_ai_for_metadata(
                        row["Keyword"],
                        row["Intent"],
                        row["Content Type"],
                        row["Notes"],
                        current_api_key,
                        ai_provider
                    )
                    
                    if output:
                        results.append({
                            "Keyword": row["Keyword"],
                            "Intent": row["Intent"],
                            "Content Type": row["Content Type"],
                            "Volume": output["volume"],
                            "SEO Title": output["seo_title"],
                            "Outline": "\n".join(output["outline"]),
                            "Instructions": output["instructions"]
                        })
                    
                    progress_bar.progress((i + 1) / len(df))
                    time.sleep(1.5)  # API rate limiting
                
                if results:
                    st.session_state["metadata_df"] = pd.DataFrame(results)
                    st.success("✅ Metadata generated successfully!")
                    st.dataframe(st.session_state["metadata_df"], use_container_width=True)
                    
                    # Download option
                    csv = st.session_state["metadata_df"].to_csv(index=False)
                    st.download_button(
                        "⬇️ Download Metadata CSV",
                        csv,
                        file_name="topic_metadata.csv",
                        mime="text/csv"
                    )
    
    elif not current_api_key:
        st.error(f"❌ {ai_provider} API key not found in secrets")

with tab2:
    st.header("📝 Step 2: Article Generation")
    
    if "metadata_df" in st.session_state and not st.session_state["metadata_df"].empty:
        metadata_df = st.session_state["metadata_df"]
        
        # Single article generation
        st.subheader("Generate Single Article")
        selected_keyword = st.selectbox("Select a topic", metadata_df["Keyword"])
        
        if selected_keyword:
            selected_row = metadata_df[metadata_df["Keyword"] == selected_keyword].iloc[0]
            
            # Allow editing of metadata
            seo_title = st.text_input("SEO Title", selected_row["SEO Title"])
            outline = st.text_area("Outline", selected_row["Outline"])
            instructions = st.text_area("Instructions", selected_row["Instructions"])
            
            if st.button("Generate Article") and current_api_key:
                with st.spinner("Generating article..."):
                    article = generate_article(
                        selected_keyword,
                        seo_title,
                        outline.split("\n"),
                        instructions,
                        current_api_key,
                        ai_provider
                    )
                    
                    if article:
                        st.session_state["articles"][selected_keyword] = article
                        st.success("✅ Article generated!")
                        st.markdown("### Preview:")
                        st.markdown(article, unsafe_allow_html=True)
        
        # Bulk generation
        st.subheader("Generate All Articles")
        if st.button("🚀 Generate All Articles") and current_api_key:
            progress_bar = st.progress(0)
            
            for i, row in metadata_df.iterrows():
                keyword = row["Keyword"]
                st.info(f"Generating article for: {keyword}")
                
                article = generate_article(
                    keyword,
                    row["SEO Title"],
                    row["Outline"].split("\n"),
                    row["Instructions"],
                    current_api_key,
                    ai_provider
                )
                
                if article:
                    st.session_state["articles"][keyword] = article
                
                progress_bar.progress((i + 1) / len(metadata_df))
                time.sleep(1.5)
            
            st.success(f"✅ Generated {len(st.session_state['articles'])} articles!")
    
    else:
        st.info("⚠️ Please complete Step 1 first to generate metadata")

with tab3:
    st.header("🔗 Step 3: Internal Linking")
    
    if st.session_state["articles"]:
        st.subheader("Upload Anchor Links Mapping")
        mapping_file = st.file_uploader(
            "Upload CSV with 'anchor,url' columns or JSON file",
            type=["csv", "json"]
        )
        
        if mapping_file:
            try:
                if mapping_file.name.endswith(".csv"):
                    link_df = pd.read_csv(mapping_file)
                    if "anchor" in link_df.columns and "url" in link_df.columns:
                        anchor_map = dict(zip(link_df["anchor"], link_df["url"]))
                        st.success(f"✅ Loaded {len(anchor_map)} anchor links")
                        st.dataframe(link_df)
                    else:
                        st.error("❌ CSV must have 'anchor' and 'url' columns")
                else:  # JSON file
                    anchor_map = json.load(mapping_file)
                    st.success(f"✅ Loaded {len(anchor_map)} anchor links from JSON")
                    st.json(anchor_map)
                
                # Apply links to articles
                if st.button("🔗 Apply Internal Links to All Articles"):
                    linked_articles = {}
                    for keyword, article in st.session_state["articles"].items():
                        linked_articles[keyword] = apply_internal_links(article, anchor_map)
                    
                    st.session_state["articles"] = linked_articles
                    st.success("✅ Internal links applied to all articles!")
                    
            except Exception as e:
                st.error(f"❌ Error processing mapping file: {str(e)}")
    
    else:
        st.info("⚠️ Please generate articles first")

with tab4:
    st.header("🖼️ Step 4: Image Generation")
    
    if st.session_state["articles"] and hf_client:
        st.subheader("Generate Images for Articles")
        
        # Single image generation
        keywords = list(st.session_state["articles"].keys())
        selected_keyword = st.selectbox("Select article for image", keywords)
        
        if selected_keyword:
            # Allow custom prompt editing
            default_prompt = f"Professional illustration for {selected_keyword}, educational content, modern design, high quality, suitable for blog post"
            image_prompt = st.text_area("Image Prompt", default_prompt)
            
            if st.button("🎨 Generate Image"):
                with st.spinner("Generating image..."):
                    image_buffer = generate_ai_image(image_prompt, hf_client)
                    
                    if image_buffer:
                        st.session_state["images"][selected_keyword] = image_buffer
                        st.success("✅ Image generated!")
                        st.image(image_buffer, caption=f"Generated for: {selected_keyword}")
        
        # Bulk image generation
        st.subheader("Generate All Images")
        if st.button("🚀 Generate Images for All Articles"):
            progress_bar = st.progress(0)
            
            for i, keyword in enumerate(keywords):
                st.info(f"Generating image for: {keyword}")
                
                prompt = f"Professional illustration for {keyword}, educational content, modern design, high quality, suitable for blog post"
                image_buffer = generate_ai_image(prompt, hf_client)
                
                if image_buffer:
                    st.session_state["images"][keyword] = image_buffer
                
                progress_bar.progress((i + 1) / len(keywords))
                time.sleep(2)  # Rate limiting
            
            st.success(f"✅ Generated {len(st.session_state['images'])} images!")
    
    elif not hf_client:
        st.error("❌ Hugging Face client not initialized. Check HF_TOKEN in secrets")
    else:
        st.info("⚠️ Please generate articles first")

with tab5:
    st.header("🚀 Step 5: WordPress Publishing")
    
    if st.session_state["articles"] and wp_base_url and wp_username and wp_password:
        wp_config = {
            "base_url": wp_base_url,
            "username": wp_username,
            "password": wp_password
        }
        
        st.subheader("Publish Settings")
        
        # Global tags for all articles
        global_tags = st.text_input("Tags (comma-separated)", "education,india,guide")
        
        # Publish mode
        publish_mode = st.radio("Publish Mode", ["Draft", "Publish Immediately"])
        publish_now = publish_mode == "Publish Immediately"
        
        # Single article publishing
        st.subheader("Publish Single Article")
        keywords = list(st.session_state["articles"].keys())
        selected_keyword = st.selectbox("Select article to publish", keywords)
        
        if selected_keyword and st.button("📤 Publish Selected Article"):
            content = st.session_state["articles"][selected_keyword]
            image_buffer = st.session_state["images"].get(selected_keyword)
            
            with st.spinner("Publishing..."):
                result = publish_to_wordpress(
                    selected_keyword,
                    content,
                    image_buffer,
                    global_tags,
                    wp_config,
                    publish_now
                )
                
                if result["success"]:
                    st.success(f"✅ Published: {result['url']}")
                    st.session_state["publish_log"].append({
                        "keyword": selected_keyword,
                        "url": result["url"],
                        "timestamp": time.strftime("%Y-%m-%d %H:%M:%S")
                    })
                else:
                    st.error(f"❌ Publishing failed: {result['error']}")
        
        # Bulk publishing
        st.subheader("Publish All Articles")
        if st.button("🚀 Publish All Articles"):
            progress_bar = st.progress(0)
            
            for i, keyword in enumerate(keywords):
                st.info(f"Publishing: {keyword}")
                
                content = st.session_state["articles"][keyword]
                image_buffer = st.session_state["images"].get(keyword)
                
                result = publish_to_wordpress(
                    keyword,
                    content,
                    image_buffer,
                    global_tags,
                    wp_config,
                    publish_now
                )
                
                if result["success"]:
                    st.success(f"✅ Published: {keyword}")
                    st.session_state["publish_log"].append({
                        "keyword": keyword,
                        "url": result["url"],
                        "timestamp": time.strftime("%Y-%m-%d %H:%M:%S")
                    })
                else:
                    st.error(f"❌ Failed: {keyword} - {result['error']}")
                
                progress_bar.progress((i + 1) / len(keywords))
                time.sleep(2)  # Rate limiting
            
            st.success("✅ Bulk publishing completed!")
    
    else:
        missing = []
        if not st.session_state["articles"]:
            missing.append("articles")
        if not wp_base_url:
            missing.append("WordPress base URL")
        if not wp_username:
            missing.append("WordPress username")
        if not wp_password:
            missing.append("WordPress password")
        
        st.error(f"❌ Missing: {', '.join(missing)}")

with tab6:
    st.header("📊 Step 6: Export & Logs")
    
    # Export articles
    if st.session_state["articles"]:
        st.subheader("Export Articles")
        
        # Individual article export
        keywords = list(st.session_state["articles"].keys())
        selected_keyword = st.selectbox("Select article to export", keywords, key="export_select")
        
        if selected_keyword:
            article_content = st.session_state["articles"][selected_keyword]
            
            # HTML export
            st.download_button(
                "⬇️ Download HTML",
                article_content,
                file_name=f"{selected_keyword.replace(' ', '_')}.html",
                mime="text/html"
            )
            
            # DOCX export
            if st.button("📄 Generate DOCX"):
                doc = Document()
                doc.add_heading(selected_keyword, 0)
                
                # Remove HTML tags for DOCX (basic cleaning)
                clean_content = re.sub(r'<[^>]+>', '', article_content)
                doc.add_paragraph(clean_content)
                
                # Save to BytesIO
                docx_buffer = BytesIO()
                doc.save(docx_buffer)
                docx_buffer.seek(0)
                
                st.download_button(
                    "⬇️ Download DOCX",
                    docx_buffer,
                    file_name=f"{selected_keyword.replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        
        # Bulk export
        st.subheader("Bulk Export")
        if st.button("📦 Create ZIP Archive"):
            zip_buffer = BytesIO()
            
            with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                for keyword, content in st.session_state["articles"].items():
                    filename = f"{keyword.replace(' ', '_')}.html"
                    zip_file.writestr(filename, content)
            
            zip_buffer.seek(0)
            
            st.download_button(
                "⬇️ Download All Articles (ZIP)",
                zip_buffer,
                file_name="seo_articles.zip",
                mime="application/zip"
            )
    
    # Publishing log
    if st.session_state["publish_log"]:
        st.subheader("Publishing Log")
        log_df = pd.DataFrame(st.session_state["publish_log"])
        st.dataframe(log_df, use_container_width=True)
        
        # Export log
        log_csv = log_df.to_csv(index=False)
        st.download_button(
            "⬇️ Download Publishing Log",
            log_csv,
            file_name="publish_log.csv",
            mime="text/csv"
        )
    
    # Reset button
    st.subheader("Reset Application")
    if st.button("🔄 Clear All Data", type="secondary"):
        st.session_state.clear()
        st.success("✅ All data cleared successfully!")
        st.rerun()

# Footer
st.markdown("---")
st.markdown("""
<div style='text-align: center; color: #666; padding: 20px;'>
    <p>📚 SEO Content Automation Pipeline | Built with Streamlit</p>
    <p>Automate your content creation workflow: Topics → Metadata → Articles → Images → WordPress</p>
</div>
""", unsafe_allow_html=True)
